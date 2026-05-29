from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "generated"
OUT.mkdir(parents=True, exist_ok=True)

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(32, 42, 52)
MUTED = RGBColor(92, 105, 116)
LIGHT = "F2F4F7"
ACCENT_FILL = "E8EEF5"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_width(table, width_dxa: int = 9360) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")


def set_font(run, size: float | None = None, bold: bool | None = None, color: RGBColor | None = None) -> None:
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def add_run(paragraph, text: str, size: float | None = None, bold: bool | None = None, color: RGBColor | None = None):
    run = paragraph.add_run(text)
    set_font(run, size=size, bold=bold, color=color)
    return run


def configure_doc(doc: Document, title: str) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 14, 6),
        ("Heading 2", 13, BLUE, 10, 4),
        ("Heading 3", 11.5, DARK_BLUE, 8, 3),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_run(header, title, 8.5, False, MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(footer, "Inventory Nexus | Confidential planning and implementation document", 8, False, MUTED)


def title_block(doc: Document, title: str, subtitle: str, audience: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    add_run(p, "INVENTORY NEXUS", 10, True, MUTED)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    add_run(p, title, 24, True, RGBColor(0, 0, 0))

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(16)
    add_run(p, subtitle, 12.5, False, MUTED)

    table = doc.add_table(rows=3, cols=2)
    table.style = "Table Grid"
    set_table_width(table)
    rows = [
        ("Audience", audience),
        ("Prepared date", date.today().strftime("%B %d, %Y")),
        ("Project status", "Verified local build with backend tests, lint, frontend build, seed import, and live API smoke test"),
    ]
    for i, (label, value) in enumerate(rows):
        table.cell(i, 0).text = label
        table.cell(i, 1).text = value
        set_cell_width(table.cell(i, 0), 1800)
        set_cell_width(table.cell(i, 1), 7560)
        set_cell_shading(table.cell(i, 0), LIGHT)
        for cell in table.rows[i].cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    set_font(run, 9.5, i == 0 and cell == table.cell(i, 0), INK)
    doc.add_paragraph()


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        add_run(p, item, 10.5, None, INK)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        add_run(p, item, 10.5, None, INK)


def add_para(doc: Document, text: str, bold_lead: str | None = None) -> None:
    p = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        add_run(p, bold_lead, 10.5, True, INK)
        add_run(p, text[len(bold_lead):], 10.5, False, INK)
    else:
        add_run(p, text, 10.5, False, INK)


def add_callout(doc: Document, title: str, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_width(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, ACCENT_FILL)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    add_run(p, title + ": ", 10.5, True, DARK_BLUE)
    add_run(p, text, 10.5, False, INK)
    doc.add_paragraph()


def add_matrix(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_width(table)
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = header
        set_cell_width(cell, widths[idx])
        set_cell_shading(cell, LIGHT)
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.space_after = Pt(0)
            for run in paragraph.runs:
                set_font(run, 9.2, True, DARK_BLUE)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
            set_cell_width(cells[idx], widths[idx])
            for paragraph in cells[idx].paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.05
                for run in paragraph.runs:
                    set_font(run, 8.8, False, INK)
    doc.add_paragraph()


def font() -> ImageFont.ImageFont:
    for candidate in [
        "/System/Library/Fonts/Supplemental/Arial.ttf",
        "/System/Library/Fonts/Helvetica.ttc",
        "/Library/Fonts/Arial.ttf",
    ]:
        if Path(candidate).exists():
            return ImageFont.truetype(candidate, 28)
    return ImageFont.load_default()


def draw_flow(path: Path, title: str, nodes: list[str], caption: str) -> None:
    img = Image.new("RGB", (1800, 760), "white")
    draw = ImageDraw.Draw(img)
    f = font()
    small = ImageFont.truetype("/System/Library/Fonts/Supplemental/Arial.ttf", 22) if Path("/System/Library/Fonts/Supplemental/Arial.ttf").exists() else f
    draw.rectangle((30, 30, 1770, 730), outline=(210, 218, 226), width=3)
    draw.text((70, 60), title, fill=(31, 77, 120), font=f)
    y = 215
    box_w = 245
    gap = 35
    x = 70
    for i, node in enumerate(nodes):
        draw.rounded_rectangle((x, y, x + box_w, y + 150), radius=18, fill=(232, 238, 245), outline=(46, 116, 181), width=3)
        words = node.split()
        lines = []
        line = ""
        for word in words:
            candidate = f"{line} {word}".strip()
            if draw.textlength(candidate, font=small) > box_w - 38:
                lines.append(line)
                line = word
            else:
                line = candidate
        if line:
            lines.append(line)
        text_y = y + 38
        for line in lines:
            draw.text((x + 20, text_y), line, fill=(32, 42, 52), font=small)
            text_y += 28
        if i < len(nodes) - 1:
            ax = x + box_w + 5
            ay = y + 75
            draw.line((ax, ay, ax + gap - 12, ay), fill=(122, 90, 0), width=4)
            draw.polygon([(ax + gap - 12, ay - 10), (ax + gap + 5, ay), (ax + gap - 12, ay + 10)], fill=(122, 90, 0))
        x += box_w + gap
    draw.text((70, 600), caption, fill=(92, 105, 116), font=small)
    img.save(path)


def report_doc() -> Path:
    doc = Document()
    configure_doc(doc, "Business Report")
    diagram = OUT / "report_architecture.png"
    draw_flow(
        diagram,
        "End to End Architecture Flow",
        ["React Dashboard", "FastAPI Service", "Domain Services", "PostgreSQL Data Store", "AI and ML Services", "Executive Insights"],
        "User actions flow through authenticated APIs into inventory, customer, analytics, and planning services. Insights return to the dashboard for operational decisions.",
    )

    title_block(
        doc,
        "Inventory Nexus Business Report",
        "Production grade inventory management, analytics, and AI planning platform",
        "Business stakeholders, client sponsors, product owners, and delivery leadership",
    )

    doc.add_heading("1. Executive Overview", level=1)
    add_para(doc, "Inventory Nexus is designed as a professional inventory management and decision intelligence platform for modern retail, wholesale, and distribution operations. The application combines operational inventory controls, dashboard reporting, customer intelligence, and AI assisted planning in one coherent system. It is not positioned as a proof of concept. The foundation has been structured so it can be demonstrated to a client today and expanded into a secure production deployment.")
    add_para(doc, "The current implementation includes a React dashboard, a FastAPI backend, a PostgreSQL ready data model, Docker based deployment, seeded product and warehouse data, and an imported customer mart containing 5,878 customer records. The platform supports authenticated access, SKU visibility, stock movement control, low stock monitoring, reorder recommendations, demand planning signals, and customer churn insight.")
    add_para(doc, "The business value is straightforward: Inventory Nexus gives decision makers a single operating view across inventory health, supplier performance, customer behavior, and replenishment risk. This reduces manual reporting effort, improves planning discipline, and creates a foundation for more advanced forecasting and recommendation use cases.")

    doc.add_heading("2. Objectives", level=1)
    add_bullets(doc, [
        "Create a client ready inventory management platform with a polished dashboard and reliable backend services.",
        "Support core inventory operations including products, suppliers, warehouses, stock balances, and auditable stock movements.",
        "Use PostgreSQL as the production data store, with Docker Compose for repeatable local and staging deployment.",
        "Integrate the customer mart dataset so business users can see customer segmentation, churn risk, and future recommendation opportunities.",
        "Define a practical AI and ML roadmap that supports demand forecasting, reorder optimization, churn prediction, and next best product recommendations.",
        "Establish a clean architecture that can grow into cloud hosting, stronger security, observability, and enterprise governance.",
    ])

    doc.add_heading("3. Technical Architecture", level=1)
    add_para(doc, "The architecture was planned around a simple principle: separate user experience, business logic, and persistence clearly enough that each layer can mature independently. The frontend is responsible for presenting decisions and workflows. The backend owns validation, authentication, data access, domain rules, analytics, and AI planning contracts. The database remains the system of record for operational truth.")
    add_para(doc, "The end to end flow begins when a business user opens the React dashboard and signs in with seeded credentials or future enterprise identity. The dashboard calls FastAPI endpoints using a bearer token. FastAPI routes delegate work to service modules that enforce inventory rules, calculate analytics, read the customer mart, and return structured responses. In production, the same service layer talks to PostgreSQL through SQLAlchemy. In local testing, SQLite is used to make fast automated checks possible.")
    add_para(doc, "This approach gives the project a strong demo experience while keeping a production path visible. The current deterministic AI planning endpoints are deliberately stable contracts. That means the dashboard does not need to change when baseline heuristics are replaced with trained forecasting, churn, and recommendation models.")

    doc.add_picture(str(diagram), width=Inches(6.45))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("Functional Modules", level=2)
    add_bullets(doc, [
        "Frontend dashboard: React and TypeScript interface for executive metrics, inventory availability, reorder queue, demand forecast, and customer segmentation.",
        "Authentication module: FastAPI login endpoint, password hashing, JWT token creation, and protected route dependencies.",
        "Inventory module: product listing, product creation, warehouse stock balances, and stock movement recording with safeguards against negative inventory.",
        "Analytics module: dashboard metrics for active SKUs, inventory value, low stock alerts, reserved units, churn risk customers, and supplier reliability.",
        "AI planning module: reorder recommendations and baseline thirty day demand forecasts exposed through stable API contracts.",
        "Customer intelligence module: imported customer mart with RFM, churn, order behavior, monetary value, region, and lifespan features.",
        "Data layer: SQLAlchemy ORM models for users, suppliers, warehouses, categories, products, stock items, stock movements, and customer segments.",
        "Deployment layer: Docker Compose stack for PostgreSQL, backend, and frontend services with environment driven configuration.",
    ])

    doc.add_heading("4. AI and ML Strategy", level=1)
    add_para(doc, "AI and ML are useful in this project because inventory management is fundamentally a prediction and optimization problem. A production team needs to know what will sell, when stock will run out, which suppliers create fulfillment risk, which customers may churn, and which products should be recommended to increase basket value. Inventory Nexus already exposes the right business surfaces for these capabilities.")
    add_para(doc, "The current application includes deterministic AI planning endpoints for immediate usability. These endpoints produce reorder recommendations from on hand stock, reorder points, supplier lead time, and reorder quantity. They also generate baseline demand forecasts from stock movement history. This is intentionally practical: the client can see the decision workflow now, while the team retains a clear path to replace baseline logic with trained models.")
    add_matrix(doc, ["AI or ML Capability", "Business Purpose", "Current Status", "Production Direction"], [
        ["Demand forecasting", "Estimate SKU level demand for the next planning window.", "Baseline endpoint available.", "Train time series or regression models using sales movements, promotions, seasonality, region, and product category."],
        ["Reorder optimization", "Recommend when and how much to reorder.", "Rule based recommendations available.", "Optimize using forecast demand, lead time variability, carrying cost, stockout cost, and supplier reliability."],
        ["Churn prediction", "Identify customers likely to stop purchasing.", "Customer mart includes churn labels and RFM features.", "Train supervised models and explain top churn drivers by segment and region."],
        ["Customer lifetime value", "Prioritize high value customers and retention spend.", "High value customer counts available.", "Build CLV bands and combine them with churn scores for targeted action."],
        ["Next best product", "Increase cross sell and repeat purchase revenue.", "Dataset supports future recommendation work.", "Use customer behavior, product affinity, category similarity, and embeddings for recommendations."],
        ["AI assistant", "Let business users ask safe operational questions.", "Architecture prepared for this path.", "Expose governed tools for inventory, customer, and analytics queries with role checks and audit logs."],
    ], [2200, 2500, 2100, 2560])

    doc.add_heading("5. Data Governance and Controls", level=1)
    add_para(doc, "Data governance is important because this platform combines operational inventory data with customer level analytics. The business should treat inventory, supplier, and customer data as controlled assets. The system design already separates domain models and API contracts, which creates a strong place to add access rules, validation, auditing, and retention policies.")
    add_bullets(doc, [
        "Data ownership: Inventory operations owns SKU, warehouse, stock movement, and supplier data. Commercial or CRM teams own customer segmentation inputs.",
        "Data quality: Product SKUs should be unique, stock movement quantities should be validated, and customer import pipelines should reject malformed or duplicate customer IDs.",
        "Auditability: Stock movements are append style operational records, which supports traceability for receipts, sales, returns, transfers, and adjustments.",
        "Privacy: Customer mart data should be handled under privacy rules, with access limited to authorized roles and reporting aggregated wherever possible.",
        "Model governance: Forecasting, churn, and recommendation models should have versioning, validation metrics, monitoring, and human override paths.",
        "Retention: Operational movement data should be retained according to finance and audit needs, while customer analytics data should follow consent and privacy policies.",
    ])

    doc.add_heading("6. Security, Compliance, and Risk View", level=1)
    add_para(doc, "The current implementation includes authenticated API access and password hashing, which is suitable for an early controlled demo. For production, the security model should evolve toward enterprise identity, role based authorization, environment managed secrets, encrypted transport, audit trails, and database level access controls.")
    add_matrix(doc, ["Area", "Current Foundation", "Recommended Next Step"], [
        ["Authentication", "JWT login with hashed password.", "Integrate OAuth or OpenID Connect with enterprise identity provider."],
        ["Authorization", "Protected API routes.", "Add role based permissions for admin, manager, analyst, and viewer access."],
        ["Secrets", "Environment based configuration.", "Move secrets to a cloud secret manager and rotate credentials."],
        ["Database", "PostgreSQL ready schema.", "Enable backups, encryption, migration discipline, and row level security where needed."],
        ["Monitoring", "Basic local verification.", "Add structured logs, metrics, tracing, uptime alerts, and audit dashboards."],
    ], [2000, 3300, 4060])

    doc.add_heading("7. Future Direction", level=1)
    add_bullets(doc, [
        "Cloud deployment on AWS, Azure, or Google Cloud with managed PostgreSQL, container hosting, private networking, and automated deployments.",
        "Database migrations through Alembic, with a formal release process for schema changes.",
        "CI pipeline for tests, linting, build validation, security scans, and Docker image publishing.",
        "Production frontend hosting through a CDN backed static hosting service.",
        "Operational modules for purchase orders, receiving, transfers, cycle counts, returns, stock reservations, and invoice reconciliation.",
        "Advanced ML services with model training jobs, model registry, drift monitoring, and explainability reports.",
        "Enterprise reporting layer for executive KPIs, supplier scorecards, inventory aging, and margin analysis.",
    ])

    doc.add_heading("8. Current Verification Status", level=1)
    add_para(doc, "The current codebase has been verified locally with locked dependency versions. Backend tests pass, backend lint passes, frontend production build passes, the seed script imports the customer mart, and live API endpoints return expected data. This gives the project a reliable baseline for continued client demonstration and technical development.")

    path = OUT / "Report.docx"
    doc.save(path)
    return path


def flow_doc() -> Path:
    doc = Document()
    configure_doc(doc, "Technical Flow Guide")
    diagram = OUT / "flow_technical.png"
    draw_flow(
        diagram,
        "Request and Code Flow",
        ["Browser UI", "API Client", "FastAPI Router", "Service Layer", "SQLAlchemy Session", "Database and Dataset"],
        "Each screen calls typed API helpers. Routers handle HTTP concerns. Services implement business rules. SQLAlchemy persists and reads the operational data model.",
    )

    title_block(
        doc,
        "Inventory Nexus Technical Flow",
        "Technology choices, codebase map, runtime flow, dependency versions, and verification guide",
        "Developers, technical reviewers, solution architects, and maintainers",
    )

    doc.add_heading("1. Purpose of This Guide", level=1)
    add_para(doc, "This guide explains how the Inventory Nexus codebase works from end to end. It is written for a reader who wants to understand the technology choices, folder structure, backend flow, frontend flow, data model, AI service contracts, and testing approach without reading every source file first.")
    add_para(doc, "The project is organized as a full stack application with a React frontend, a FastAPI backend, a PostgreSQL ready persistence layer, and Docker Compose orchestration. The same architecture supports local development, client demonstrations, and a future cloud deployment.")

    doc.add_heading("2. Repository Layout", level=1)
    add_matrix(doc, ["Path", "Purpose"], [
        ["backend/app/main.py", "Creates the FastAPI app, configures CORS, creates database tables for the current baseline, and registers API routes."],
        ["backend/app/api/v1", "HTTP route layer for authentication, inventory, and analytics endpoints."],
        ["backend/app/core", "Configuration and security helpers, including environment settings, password hashing, JWT creation, and token decoding."],
        ["backend/app/db/session.py", "SQLAlchemy engine, declarative base, session factory, and database dependency."],
        ["backend/app/models/domain.py", "SQLAlchemy domain models for users, suppliers, warehouses, categories, products, stock items, stock movements, and customer segments."],
        ["backend/app/schemas/domain.py", "Pydantic request and response models used by API endpoints."],
        ["backend/app/services", "Business logic for inventory, analytics, ML style planning, and seed data import."],
        ["backend/tests/test_api.py", "API tests that validate health, login, seeded products, dashboard data, and reorder recommendations."],
        ["frontend/src/pages/App.tsx", "Main dashboard page. It signs in, loads API data, and renders the operating view."],
        ["frontend/src/lib/api.ts", "Typed frontend API helper functions and TypeScript data shapes."],
        ["docker-compose.yml", "Local container stack for PostgreSQL, backend, and frontend."],
    ], [2800, 6560])

    doc.add_heading("3. Technology Choices", level=1)
    add_bullets(doc, [
        "React with TypeScript was selected for the dashboard because it provides a maintainable component model, static typing, and strong ecosystem support for business applications.",
        "Vite was selected for frontend tooling because it gives fast local development and a simple production build path.",
        "FastAPI was selected for the backend because it provides high performance, automatic OpenAPI documentation, dependency injection, and native Pydantic validation.",
        "SQLAlchemy was selected for persistence because it offers mature ORM mapping, explicit session control, and portability across SQLite testing and PostgreSQL production.",
        "PostgreSQL is the intended production database because inventory and customer analytics need transactional reliability, indexing, relational integrity, and future support for advanced extensions.",
        "Docker Compose was selected for repeatable local and staging runs across frontend, backend, and database services.",
    ])

    doc.add_heading("4. End to End Runtime Flow", level=1)
    doc.add_picture(str(diagram), width=Inches(6.45))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_numbered(doc, [
        "The user opens the dashboard in the browser. The dashboard attempts to sign in with the seeded demo account for local demonstration.",
        "The frontend sends a login request to the FastAPI authentication endpoint and stores the returned bearer token in local storage.",
        "The dashboard requests metrics, products, reorder recommendations, forecasts, and customer insights using the bearer token.",
        "FastAPI validates the token through the shared dependency in backend/app/api/v1/deps.py.",
        "Routers call service functions. The service layer owns business rules, such as preventing stock from becoming negative and calculating low stock recommendations.",
        "The SQLAlchemy session reads and writes records through the models in backend/app/models/domain.py.",
        "The response returns as structured JSON. The React dashboard renders metric cards, tables, alerts, forecasts, and customer summaries.",
    ])

    doc.add_heading("5. Backend Flow in Detail", level=1)
    add_para(doc, "The backend starts in backend/app/main.py. When imported, it loads settings, creates database tables for the baseline implementation, configures CORS, and registers the v1 router under /api/v1. The health endpoint remains outside the versioned router so infrastructure can check service availability quickly.")
    add_para(doc, "Route files stay thin. The auth route validates credentials and returns a token. The inventory route lists products, creates products, and records stock movements. The analytics route returns dashboard metrics, reorder recommendations, demand forecasts, and customer insights. This keeps HTTP details separate from business behavior.")
    add_para(doc, "The service layer is where the application becomes more than CRUD. Inventory services aggregate stock quantities, enforce SKU uniqueness, apply movement deltas, stop negative stock, calculate inventory value, and produce reorder recommendations. ML services expose forecast and customer insight contracts. Seed services create demo users, suppliers, warehouses, products, stock, and customer segment records.")

    doc.add_heading("6. Data Model Summary", level=1)
    add_matrix(doc, ["Entity", "What it Represents", "Important Relationships"], [
        ["User", "Authenticated application user.", "Used by login and token protected routes."],
        ["Supplier", "Vendor that provides products.", "Products reference suppliers; supplier lead time and reliability influence reorder logic."],
        ["Warehouse", "Physical or logical stock location.", "Stock items and stock movements are tied to a warehouse."],
        ["Category", "Product grouping.", "Products belong to categories for reporting and filtering."],
        ["Product", "Sellable or stocked SKU.", "Connected to category, supplier, stock items, and stock movements."],
        ["StockItem", "Current on hand and reserved stock by product and warehouse.", "Aggregated for product availability and low stock checks."],
        ["StockMovement", "Auditable inventory event.", "Records receipts, sales, adjustments, transfers, and returns."],
        ["CustomerSegment", "Imported customer mart record.", "Supports churn, RFM, regional analysis, CLV, and future recommendations."],
    ], [1700, 3500, 4160])

    doc.add_heading("7. Frontend Flow in Detail", level=1)
    add_para(doc, "The frontend entry point is frontend/src/main.tsx. It renders App.tsx and loads the shared stylesheet. App.tsx owns the dashboard state, signs in through the API helper, loads five API resources in parallel, and renders the operational dashboard.")
    add_bullets(doc, [
        "MetricCard renders compact KPI tiles for active SKUs, inventory value, low stock alerts, and churn risk.",
        "DataTable renders product availability, category, supplier, stock position, and margin.",
        "The reorder queue renders recommendations from the analytics API and explains why each product needs action.",
        "The forecast panel renders expected thirty day demand and confidence values.",
        "The customer mart panel renders total imported customers, churn rate, high value customers, and top regions.",
    ])

    doc.add_heading("8. AI and ML Flow", level=1)
    add_para(doc, "The current AI and ML implementation is intentionally practical. The application exposes stable planning endpoints today and can replace the internal logic with trained models later. This allows frontend, backend, and business stakeholders to agree on workflow before the team invests in heavier model training infrastructure.")
    add_bullets(doc, [
        "Reorder recommendations use current stock, reorder point, reorder quantity, supplier lead time, and urgency classification.",
        "Demand forecast uses stock movement history as a baseline and returns expected thirty day demand, confidence, and method.",
        "Customer insights use the customer mart to calculate total customers, churn rate, high value customers, and leading regions.",
        "Future models can use the same endpoint contracts, which reduces frontend rework and keeps adoption smoother.",
    ])

    doc.add_heading("9. Dependency Versions", level=1)
    add_para(doc, "The backend dependency file is fully pinned to reduce testing drift. Important versions are listed below. The complete source of truth is backend/requirements.txt.")
    add_matrix(doc, ["Dependency", "Pinned Version", "Role"], [
        ["fastapi", "0.115.6", "Backend web framework and OpenAPI documentation."],
        ["uvicorn", "0.34.0", "ASGI server for local and container runtime."],
        ["SQLAlchemy", "2.0.50", "ORM and database session layer."],
        ["psycopg and psycopg-binary", "3.2.13", "PostgreSQL driver."],
        ["pydantic", "2.13.4", "Data validation and response schemas."],
        ["pydantic-settings", "2.7.1", "Environment driven configuration."],
        ["python-jose", "3.3.0", "JWT creation and validation."],
        ["passlib", "1.7.4", "Password hashing abstraction."],
        ["bcrypt", "4.0.1", "Password hashing backend compatible with Passlib 1.7.4."],
        ["pytest", "8.3.4", "Backend test runner."],
        ["httpx", "0.28.1", "Test client and smoke test HTTP calls."],
        ["ruff", "0.8.4", "Backend linting."],
    ], [2600, 1900, 4860])

    add_para(doc, "The frontend top level dependencies are pinned in frontend/package.json and the resolved dependency tree is captured in frontend/package-lock.json. Verified top level versions include React 19.2.6, React DOM 19.2.6, Vite 6.4.2, TypeScript 5.9.3, lucide-react 0.468.0, @vitejs/plugin-react 4.7.0, @types/react 19.2.15, and @types/react-dom 19.2.3.")

    doc.add_heading("10. Local Commands for Understanding and Verification", level=1)
    add_bullets(doc, [
        "Create or activate the backend environment: cd backend, then source .venv/bin/activate.",
        "Install locked backend dependencies: pip install -r requirements.txt.",
        "Seed local data: python scripts/seed.py.",
        "Run backend tests: pytest -q.",
        "Run backend lint: ruff check app tests scripts.",
        "Start the backend locally: uvicorn app.main:app --host 127.0.0.1 --port 8000.",
        "Install frontend dependencies: cd frontend, then npm install.",
        "Build the frontend: npm run build.",
        "Run the full container stack: docker compose up --build, then docker compose exec backend python scripts/seed.py.",
    ])

    doc.add_heading("11. Verification Results", level=1)
    add_bullets(doc, [
        "Backend tests passed with three API tests covering health, authentication, seeded dashboard data, product listing, and reorder recommendations.",
        "Backend lint passed with Ruff across app, tests, and scripts.",
        "Frontend production build passed with Vite and TypeScript.",
        "Seed script successfully created schema and imported the 5,878 row customer mart.",
        "Live API smoke test returned successful responses for health, login, dashboard, products, reorder recommendations, and customer insights.",
    ])

    doc.add_heading("12. Technical Improvement Backlog", level=1)
    add_bullets(doc, [
        "Add Alembic migration files instead of relying on create_all during application startup.",
        "Add role based permissions and replace demo login with enterprise identity integration.",
        "Add integration tests against PostgreSQL in Docker for closer production parity.",
        "Add structured logging, request IDs, metrics, and traces.",
        "Add CI workflow for backend tests, frontend builds, linting, dependency review, and Docker image checks.",
        "Add model training jobs, model registry, and monitoring for the AI and ML roadmap.",
    ])

    path = OUT / "Flow.docx"
    doc.save(path)
    return path


if __name__ == "__main__":
    print(report_doc())
    print(flow_doc())
